"""
export_docx.py — Genera il documento Word "Allegato Documenti di Progetto"
partendo dai dati di un ProjectSizing salvato.
"""
import io
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Colori brand ──────────────────────────────────────────────────────────────
BLUE_DARK   = RGBColor(0x1D, 0x3A, 0x57)   # intestazioni scure
BLUE_ACCENT = RGBColor(0x1D, 0x6E, 0xF5)   # accento
GRAY_TEXT   = RGBColor(0x6B, 0x72, 0x80)   # testo secondario
GRAY_LIGHT  = RGBColor(0xF0, 0xF4, 0xFF)   # sfondo header tabella
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
BLACK       = RGBColor(0x1A, 0x1D, 0x23)

TAGLIA_INFO = {
    "xs":  {"label": "XS",  "range": "< 2.000 €"},
    "s":   {"label": "S",   "range": "2.000 – 3.000 €"},
    "m":   {"label": "M",   "range": "5.000 – 8.000 €"},
    "l":   {"label": "L",   "range": "9.000 – 15.000 €"},
    "xl":  {"label": "XL",  "range": "16.000 – 30.000 €"},
    "xxl": {"label": "XXL", "range": "> 30.000 €"},
}

TAGLIA_ORDER = ["xs", "s", "m", "l", "xl", "xxl"]

def _ti(t): return TAGLIA_ORDER.index(t) if t in TAGLIA_ORDER else -1

Q_LABELS = {
    "q1":  {"label": "Giornate di sviluppo",      "opts": {"0": "< 2 gg", "1": "2–10 gg", "3": "10–50 gg", "4": "> 50 gg"}},
    "q2":  {"label": "Durata",                    "opts": {"0": "< 1 mese", "1": "1–3 mesi", "3": "3–6 mesi", "4": "> 6 mesi"}},
    "q3":  {"label": "Team interni",              "opts": {"0": "1 team", "1": "2 team", "2": "3+ team"}},
    "q4":  {"label": "Vendor interni",            "opts": {"0": "Nessuno", "1": "1 vendor", "2": "2+ vendor"}},
    "q5":  {"label": "Vendor esterni",            "opts": {"0": "Nessuno", "1.5": "1 vendor", "2.5": "2+ vendor"}},
    "q6":  {"label": "Referenti funzionali",      "opts": {"0": "1–2 persone", "1": "3–5 persone", "2": "6+ persone"}},
    "q7":  {"label": "Utenti finali impattati",   "opts": {"0": "Nessuno", "1": "1–5 utenti", "2": "> 5 utenti"}},
    "q8":  {"label": "SAP coinvolto",             "opts": {"0": "No", "2": "Sì"}},
    "q9":  {"label": "Vincoli compliance",        "opts": {"0": "No", "2": "Sì"}},
    "q10": {"label": "Go-live tecnico",           "opts": {"0": "No", "1": "Sì"}},
    "q11": {"label": "Cliente top 20",            "opts": {"0": "No", "1": "Sì"}},
}

GOVERNANCE_ITEMS = [
    # (name, tipo, fase, from_taglia, frequenza, descrizione, cond_fn or None)
    ("Verbali e action item log",       "D", "Esecuzione",    "xs", "Ad ogni riunione",
     "Registro delle decisioni prese e delle azioni assegnate durante le riunioni di progetto, con responsabile e scadenza.",
     None),
    ("Piano di progetto (baseline)",    "D", "Pianificazione", "s",  "Una tantum — aggiornato a ogni variante approvata",
     "Documento che formalizza scope, milestones, timeline, risorse e dipendenze del progetto nella versione approvata.",
     None),
    ("Kick-off",                        "A", "Avvio",          "s",  "Una tantum",
     "Riunione iniziale con il team di progetto e i referenti cliente per allineare obiettivi, scope, ruoli e modalità operative.",
     None),
    ("Milestone plan",                  "D", "Pianificazione", "m",  "Una tantum — condiviso con Steering Committee",
     "Piano sintetico dei principali punti di controllo del progetto con le date attese di completamento.",
     None),
    ("RACI",                            "D", "Avvio",          "m",  "Una tantum — rivisto a ogni cambio scope",
     "Matrice che definisce chi è Responsabile, Approvatore, Consultato e Informato per ciascuna attività di progetto.",
     None),
    ("Risk Register",                   "D", "Pianificazione", "m",  "Continuo (documento vivo)",
     "Documento vivo che censisce i rischi identificati, la loro probabilità, impatto e le azioni di mitigazione associate.",
     None),
    ("Issue Log",                       "D", "Esecuzione",     "m",  "Continuo (documento vivo)",
     "Registro continuo dei problemi aperti che impattano l'esecuzione del progetto, con stato e responsabile.",
     None),
    ("Meeting operativi strutturati",   "A", "Esecuzione",     "m",  "Settimanale",
     "Riunioni periodiche con agenda predefinita per monitorare avanzamento, blocchi e azioni in corso.",
     None),
    ("Organigramma di progetto",        "D", "Avvio",          "m",  "Una tantum — aggiornato a ogni cambio ruolo",
     "Schema visivo delle figure coinvolte nel progetto con relativi ruoli e linee di riporto.",
     lambda a, t: _ti(t) >= _ti("l") or (_ti(t) == _ti("m") and (float(a.get("q4", 0)) > 0 or float(a.get("q5", 0)) > 0))),
    ("Stakeholder Register",            "D", "Avvio",          "l",  "Una tantum — aggiornato se cambiano gli attori",
     "Registro degli stakeholder chiave con ruolo, livello di influenza, aspettative e strategia di coinvolgimento.",
     None),
    ("Piano di comunicazione",          "D", "Avvio",          "l",  "Una tantum",
     "Documento che definisce cosa comunicare, a chi, con quale frequenza e attraverso quale canale.",
     None),
    ("Change management process",       "A", "Esecuzione",     "l",  "On demand",
     "Processo strutturato per valutare, approvare e gestire le richieste di variazione di scope, budget o timeline.",
     None),
    ("Escalation process",              "A", "Esecuzione",     "l",  "On demand",
     "Procedura che definisce come e quando un problema viene portato a un livello decisionale superiore.",
     None),
    ("Gate review / Phase gate",        "A", "Fine di ogni fase", "l", "Una tantum per fase",
     "Checkpoint formale al termine di ogni fase per validare i deliverable e autorizzare il passaggio alla fase successiva.",
     None),
    ("Steering Committee",              "A", "Esecuzione → Chiusura", "l", "Mensile / bimestrale",
     "Incontro periodico con il management di Archiva e del cliente per monitorare l'avanzamento strategico.",
     lambda a, t: _ti(t) >= _ti("xl") or (_ti(t) == _ti("l") and float(a.get("q11", 0)) > 0)),
    ("Alignment Report",                "D", "Esecuzione",     "l",  "Settimanale (Confluence)",
     "Report minimalista pubblicato settimanalmente su Confluence: % avanzamento, attività in corso, problematiche e prossime attività. Il commerciale viene taggato ad ogni aggiornamento.",
     None),
    ("Governance Effort Report",        "D", "Esecuzione → Chiusura", "l", "Mensile / a fine progetto",
     "Report interno che rendiconta le ore spese in attività di governance. Prodotto mensilmente e a consuntivo a fine progetto.",
     None),
    ("Vendor management plan",          "D", "Pianificazione", "xl", "Una tantum — aggiornato a nuovi contratti",
     "Piano che definisce modalità di gestione dei fornitori: SLA attesi, interfacce operative e processi di escalation.",
     None),
    ("Cutover Plan",                    "D", "Pre go-live",    "xl", "Una tantum — aggiornato fino al go-live",
     "Documento che dettaglia la sequenza di attività, responsabili e timing per il passaggio in produzione del sistema.",
     lambda a, t: _ti(t) >= _ti("xl") or (_ti(t) == _ti("l") and float(a.get("q10", 0)) > 0)),
    ("Rollback Plan",                   "D", "Pre go-live",    "xl", "Una tantum — validato prima del go-live",
     "Procedura di rientro in caso di blocco critico durante il cutover, con trigger, responsabili e tempi massimi.",
     lambda a, t: _ti(t) >= _ti("xl") or (_ti(t) == _ti("l") and float(a.get("q10", 0)) > 0)),
    ("Hypercare",                       "A", "Post go-live",   "xl", "Settimanale · 4–8 settimane",
     "Periodo di sorveglianza intensiva post go-live con monitoraggio settimanale e supporto prioritario.",
     lambda a, t: _ti(t) >= _ti("xl") or (t in ("m", "l") and float(a.get("q10", 0)) > 0 and float(a.get("q7", 0)) >= 2)),
    ("Handover document",               "D", "Chiusura",       "xl", "Una tantum — con sign-off del ricevente",
     "Documento formale che formalizza il passaggio del progetto in gestione operativa al team di operations.",
     None),
    ("Handover a operations",           "A", "Go-live / Chiusura", "xl", "Una tantum",
     "Attività strutturata di trasferimento di conoscenza e responsabilità dal team di progetto al team di operations.",
     None),
]


def get_active_items(answers: dict, taglia: str, gov_map: dict = None):
    """Restituisce la lista degli elementi di governance attivi per la taglia e le risposte date."""
    active = []
    ti = _ti(taglia)
    for item in GOVERNANCE_ITEMS:
        name, tipo, fase, from_t, freq, desc, cond = item
        fi = _ti(from_t)
        # Override with DB values if available
        if gov_map:
            for db_item in gov_map.values():
                if db_item.get("name") == name or db_item.get("name", "").lower() == name.lower():
                    name     = db_item.get("name", name)
                    freq     = db_item.get("frequenza", freq)
                    desc     = db_item.get("descrizione", desc)
                    break
        if cond is not None:
            if cond(answers, taglia):
                conditional = ti < fi
                active.append((name, tipo, fase, freq, desc, conditional))
        else:
            if ti >= fi:
                active.append((name, tipo, fase, freq, desc, False))
    return active


# ── Helpers XML / formatting ───────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_borders(cell, color="DADDE8"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _set_no_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _para(text, bold=False, size=11, color=BLACK, align=WD_ALIGN_PARAGRAPH.LEFT, italic=False, font="Arial"):
    from docx.oxml import OxmlElement as OE
    p = OE("w:p")
    pPr = OE("w:pPr")
    jc = OE("w:jc")
    jc.set(qn("w:val"), {
        WD_ALIGN_PARAGRAPH.LEFT: "left",
        WD_ALIGN_PARAGRAPH.CENTER: "center",
        WD_ALIGN_PARAGRAPH.RIGHT: "right",
    }.get(align, "left"))
    pPr.append(jc)
    p.append(pPr)
    if text:
        r = OE("w:r")
        rPr = OE("w:rPr")
        rFonts = OE("w:rFonts")
        rFonts.set(qn("w:ascii"), font)
        rFonts.set(qn("w:hAnsi"), font)
        rPr.append(rFonts)
        sz = OE("w:sz"); sz.set(qn("w:val"), str(size * 2)); rPr.append(sz)
        szCs = OE("w:szCs"); szCs.set(qn("w:val"), str(size * 2)); rPr.append(szCs)
        if bold:
            rPr.append(OE("w:b"))
        if italic:
            rPr.append(OE("w:i"))
        clr = OE("w:color")
        clr.set(qn("w:val"), f"{color[0]:02X}{color[1]:02X}{color[2]:02X}" if isinstance(color, tuple) else str(color).lstrip("#"))
        rPr.append(clr)
        r.append(rPr)
        t = OE("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = text
        r.append(t)
        p.append(r)
    return p


def _add_para(doc_or_cell, text, bold=False, size=11, color=None, align=WD_ALIGN_PARAGRAPH.LEFT, italic=False, space_before=0, space_after=6):
    if color is None:
        color = BLACK
    p = doc_or_cell.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p


def _set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)


# ── Main generator ─────────────────────────────────────────────────────────────

def generate_sizing_docx(sizing_data: dict, gov_map: dict = None) -> bytes:
    """
    sizing_data keys: id, crm_number, odl, cliente, titolo, answers,
                      score, taglia, note, created_at, created_by, created_by_nome
    """
    answers  = sizing_data.get("answers", {})
    taglia   = sizing_data.get("taglia", "xs")
    cliente  = sizing_data.get("cliente", "")
    titolo   = sizing_data.get("titolo", "")
    crm      = sizing_data.get("crm_number") or "—"
    odl      = sizing_data.get("odl") or "—"
    score    = sizing_data.get("score", 0)
    pm_nome  = sizing_data.get("created_by_nome", "")
    note     = sizing_data.get("note") or ""
    t_info   = TAGLIA_INFO.get(taglia, {"label": taglia.upper(), "range": ""})

    raw_date = sizing_data.get("created_at", "")
    try:
        dt = datetime.fromisoformat(raw_date.replace("Z", "+00:00"))
        data_str = dt.strftime("%d/%m/%Y")
    except Exception:
        data_str = datetime.now().strftime("%d/%m/%Y")

    active_items = get_active_items(answers, taglia, gov_map)

    doc = Document()

    # ── Page setup: A4 portrait ───────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)

    # ── Default style ─────────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)

    # ══════════════════════════════════════════════════════════════════════════
    # COPERTINA
    # ══════════════════════════════════════════════════════════════════════════

    # Spazio superiore
    for _ in range(6):
        _add_para(doc, "", size=11, space_before=0, space_after=0)

    # Linea separatrice
    p_line = doc.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_line.paragraph_format.space_before = Pt(0)
    p_line.paragraph_format.space_after = Pt(32)
    pPr = p_line._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1D6EF5")
    pBdr.append(bottom)
    pPr.append(pBdr)

    # Titolo copertina
    _add_para(doc, "Allegato Documenti di Progetto", bold=True, size=22,
              color=BLUE_DARK, align=WD_ALIGN_PARAGRAPH.CENTER,
              space_before=0, space_after=16)

    # ODL
    odl_display = f"ODL — {odl}" if odl != "—" else "ODL — N/D"
    _add_para(doc, odl_display, bold=True, size=16,
              color=BLUE_ACCENT, align=WD_ALIGN_PARAGRAPH.CENTER,
              space_before=0, space_after=24)

    # Cliente
    _add_para(doc, cliente, bold=True, size=18,
              color=BLUE_DARK, align=WD_ALIGN_PARAGRAPH.CENTER,
              space_before=0, space_after=8)

    # Titolo progetto
    _add_para(doc, titolo, bold=False, size=13,
              color=GRAY_TEXT, align=WD_ALIGN_PARAGRAPH.CENTER,
              space_before=0, space_after=48)

    # Footer copertina
    _add_para(doc, f"Data: {data_str}  ·  PM: {pm_nome}",
              size=10, color=GRAY_TEXT,
              align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=0)

    # ── Page break ────────────────────────────────────────────────────────────
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # INTESTAZIONE PROGETTO
    # ══════════════════════════════════════════════════════════════════════════

    _add_para(doc, "Dati Progetto", bold=True, size=13,
              color=BLUE_DARK, space_before=0, space_after=8)

    # Tabella info progetto (2 colonne)
    info_rows = [
        ("Cliente",          cliente),
        ("Titolo progetto",  titolo),
        ("N° CRM",           crm),
        ("ODL",              odl),
        ("Project Manager",  pm_nome),
        ("Data sizing",      data_str),
    ]
    if note:
        info_rows.append(("Note", note))

    tbl_info = doc.add_table(rows=len(info_rows), cols=2)
    tbl_info.style = "Table Grid"

    col_w = [4.5, 11.5]
    for i, (k, v) in enumerate(info_rows):
        row = tbl_info.rows[i]
        kc, vc = row.cells[0], row.cells[1]
        kc.width = Cm(col_w[0])
        vc.width = Cm(col_w[1])
        _set_cell_bg(kc, "EFF4FF")
        _set_cell_bg(vc, "FFFFFF")
        _set_cell_borders(kc)
        _set_cell_borders(vc)

        kp = kc.paragraphs[0]
        kr = kp.add_run(k)
        kr.bold = True; kr.font.name = "Arial"; kr.font.size = Pt(9.5)
        kr.font.color.rgb = BLUE_DARK
        kp.paragraph_format.space_before = Pt(2)
        kp.paragraph_format.space_after  = Pt(2)

        vp = vc.paragraphs[0]
        vr = vp.add_run(v)
        vr.font.name = "Arial"; vr.font.size = Pt(9.5)
        vr.font.color.rgb = BLACK
        vp.paragraph_format.space_before = Pt(2)
        vp.paragraph_format.space_after  = Pt(2)

    _add_para(doc, "", size=6, space_before=0, space_after=0)

    # Riepilogo parametri
    _add_para(doc, "Parametri di sizing", bold=True, size=13,
              color=BLUE_DARK, space_before=12, space_after=8)

    qdefs = ["q1", "q2", "q3", "q4", "q5", "q6", "q7", "q8", "q9", "q10", "q11"]
    tbl_q = doc.add_table(rows=len(qdefs) + 1, cols=3)
    tbl_q.style = "Table Grid"

    # Header
    hrow = tbl_q.rows[0]
    for ci, txt in enumerate(["Parametro", "Risposta", "Punti"]):
        c = hrow.cells[ci]
        _set_cell_bg(c, "1D3A57")
        _set_cell_borders(c, "1D3A57")
        p2 = c.paragraphs[0]
        r2 = p2.add_run(txt)
        r2.bold = True; r2.font.name = "Arial"; r2.font.size = Pt(9)
        r2.font.color.rgb = WHITE
        p2.paragraph_format.space_before = Pt(2)
        p2.paragraph_format.space_after  = Pt(2)

    for i, qkey in enumerate(qdefs):
        qd = Q_LABELS.get(qkey, {})
        val = answers.get(qkey, "—")
        label = qd.get("opts", {}).get(str(val), str(val))
        pts   = float(val) if val not in ("—", "") else 0
        bg = "F8FAFF" if i % 2 == 0 else "FFFFFF"
        row = tbl_q.rows[i + 1]
        widths = [5.5, 8.5, 2]
        for ci, (txt, w) in enumerate(zip([qd.get("label", qkey), label, str(pts).rstrip("0").rstrip(".")], widths)):
            c = row.cells[ci]
            c.width = Cm(w)
            _set_cell_bg(c, bg)
            _set_cell_borders(c)
            p3 = c.paragraphs[0]
            r3 = p3.add_run(txt)
            r3.font.name = "Arial"; r3.font.size = Pt(9)
            r3.font.color.rgb = GRAY_TEXT if ci == 0 else BLACK
            r3.bold = (ci == 2)
            p3.paragraph_format.space_before = Pt(2)
            p3.paragraph_format.space_after  = Pt(2)

    # ── Page break ────────────────────────────────────────────────────────────
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # LISTA GOVERNANCE
    # ══════════════════════════════════════════════════════════════════════════

    _add_para(doc, "Elementi di Governance Attivati", bold=True, size=13,
              color=BLUE_DARK, space_before=0, space_after=4)

    _add_para(doc,
              f"Per il progetto «{titolo}» (taglia {t_info['label']}) sono stati attivati "
              f"{len(active_items)} elementi di governance, di cui "
              f"{sum(1 for x in active_items if x[1]=='D')} deliverable e "
              f"{sum(1 for x in active_items if x[1]=='A')} attività.",
              size=10, color=GRAY_TEXT, space_before=0, space_after=12)

    # Tabella governance
    hdr_cols = ["Elemento", "T", "Fase", "Frequenza", "Descrizione"]
    col_widths_gov = [3.8, 0.6, 2.6, 3.5, 5.5]
    n_rows = len(active_items) + 1

    tbl_gov = doc.add_table(rows=n_rows, cols=5)
    tbl_gov.style = "Table Grid"

    # Header row
    hr = tbl_gov.rows[0]
    for ci, (h, w) in enumerate(zip(hdr_cols, col_widths_gov)):
        c = hr.cells[ci]
        c.width = Cm(w)
        _set_cell_bg(c, "1D3A57")
        _set_cell_borders(c, "1D3A57")
        p4 = c.paragraphs[0]
        r4 = p4.add_run(h)
        r4.bold = True; r4.font.name = "Arial"; r4.font.size = Pt(9)
        r4.font.color.rgb = WHITE
        p4.paragraph_format.space_before = Pt(2)
        p4.paragraph_format.space_after  = Pt(2)
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci == 1 else WD_ALIGN_PARAGRAPH.LEFT

    # Data rows
    for ri, (name, tipo, fase, freq, desc, conditional) in enumerate(active_items):
        row = tbl_gov.rows[ri + 1]
        bg = "F8FAFF" if ri % 2 == 0 else "FFFFFF"
        tipo_bg = "DBEAFE" if tipo == "D" else "FCE7F3"
        tipo_color = RGBColor(0x1D, 0x4E, 0xD8) if tipo == "D" else RGBColor(0xBE, 0x18, 0x5D)

        vals = [name, tipo, fase, freq, desc]
        for ci, (txt, w) in enumerate(zip(vals, col_widths_gov)):
            c = row.cells[ci]
            c.width = Cm(w)
            _set_cell_bg(c, tipo_bg if ci == 1 else bg)
            _set_cell_borders(c)
            p5 = c.paragraphs[0]
            r5 = p5.add_run(txt)
            r5.font.name = "Arial"
            r5.font.size = Pt(8.5 if ci == 4 else 9)
            r5.bold = (ci == 0)
            r5.font.color.rgb = tipo_color if ci == 1 else (GRAY_TEXT if ci in (2, 3) else BLACK)
            p5.paragraph_format.space_before = Pt(2)
            p5.paragraph_format.space_after  = Pt(2)
            p5.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci == 1 else WD_ALIGN_PARAGRAPH.LEFT

        # Nota condizionale
        if conditional:
            row.cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x7C, 0x3A, 0xED)
            note_p = row.cells[4].add_paragraph()
            nr = note_p.add_run("⚑ Elemento condizionale — verificare i prerequisiti")
            nr.font.name = "Arial"; nr.font.size = Pt(7.5)
            nr.font.color.rgb = RGBColor(0x7C, 0x3A, 0xED)
            nr.italic = True
            note_p.paragraph_format.space_before = Pt(1)
            note_p.paragraph_format.space_after  = Pt(1)

    # Legenda
    _add_para(doc, "", size=6, space_before=0, space_after=0)
    _add_para(doc, "D = Deliverable (documento prodotto e mantenuto)   ·   A = Attività o processo   ·   ⚑ = attivo solo al verificarsi della condizione indicata",
              size=8, color=GRAY_TEXT, italic=True, space_before=4, space_after=0)

    # ── Salva in buffer ───────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
